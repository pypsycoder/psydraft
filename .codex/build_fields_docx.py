import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(95, 99, 104)


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tcpr.append(shd)


def margins(cell, top=70, start=90, bottom=70, end=90):
    tcpr = cell._tc.get_or_add_tcPr()
    tcMar = tcpr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcpr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}")) or OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None: tcMar.append(node)


def add_inline(paragraph, text, size=10.5, color=None):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    for part in pattern.split(text):
        if not part: continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); set_font(run, "Consolas", max(8.0, size-1), color=color)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_font(run, size=size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1]); set_font(run, size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(part); set_font(run, size=size, color=color)


def split_row(line):
    s = line.strip().strip("|")
    return [c.strip().replace("\\|", "|") for c in re.split(r"(?<!\\)\|", s)]


def is_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}", line))


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trpr.append(el)


def set_cell_width(cell, dxa):
    tcpr = cell._tc.get_or_add_tcPr(); tcw = tcpr.find(qn("w:tcW"))
    if tcw is None: tcw = OxmlElement("w:tcW"); tcpr.append(tcw)
    tcw.set(qn("w:w"), str(dxa)); tcw.set(qn("w:type"), "dxa")


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = False; table.style = "Table Grid"
    total = 9360
    lengths = [max(6, max((len(r[i]) if i < len(r) else 0) for r in rows)) for i in range(cols)]
    weights = [min(42, max(10, x)) for x in lengths]
    widths = [max(720, int(total*w/sum(weights))) for w in weights]
    widths[-1] += total - sum(widths)
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW")); tblw.set(qn("w:w"), str(total)); tblw.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblpr.append(layout)
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblpr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for ri, data in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci); set_cell_width(cell, widths[ci]); margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            add_inline(p, data[ci] if ci < len(data) else "", 7.1 if cols >= 5 else 7.8)
            if ri == 0:
                shade(cell, "E8EEF5")
                for run in p.runs: run.bold = True
        if ri == 0: set_repeat_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure(doc, title):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.left_margin = sec.right_margin = Inches(1.0)
    sec.header_distance = sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (("Title",24,DARK_BLUE,0,10),("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",11.5,DARK_BLUE,10,5)):
        s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=True
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    hdr=sec.header.paragraphs[0]; hdr.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hdr.add_run(title), size=8, color=GRAY)
    foot=sec.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); foot._p.append(fld)


def convert(src, dst):
    text = Path(src).read_text(encoding="utf-8-sig")
    lines = text.splitlines(); doc = Document(); configure(doc, Path(src).stem)
    i=0
    while i < len(lines):
        line=lines[i]
        if not line.strip() or line.strip()=="---": i+=1; continue
        m=re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level=min(len(m.group(1)),3); p=doc.add_paragraph(style="Title" if level==1 and i<3 else f"Heading {level}")
            add_inline(p,m.group(2), 24 if level==1 and i<3 else (16 if level==1 else 13 if level==2 else 11.5))
            i+=1; continue
        if line.lstrip().startswith("|") and i+1<len(lines) and is_separator(lines[i+1]):
            rows=[split_row(line)]; i+=2
            while i<len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(split_row(lines[i])); i+=1
            add_table(doc,rows); continue
        lm=re.match(r"^\s*(?:[-*+]|(\d+)\.)\s+(.*)$",line)
        if lm:
            p=doc.add_paragraph(style="List Number" if lm.group(1) else "List Bullet"); p.paragraph_format.space_after=Pt(3)
            add_inline(p,lm.group(2)); i+=1; continue
        parts=[line.strip()]; i+=1
        while i<len(lines) and lines[i].strip() and not re.match(r"^(#{1,6})\s+|^\s*\||^\s*(?:[-*+]|\d+\.)\s+|^---$",lines[i]):
            parts.append(lines[i].strip()); i+=1
        p=doc.add_paragraph(); add_inline(p," ".join(parts))
    doc.core_properties.title = Path(src).stem
    doc.core_properties.author = "PsyDraft"
    doc.save(dst)


if __name__ == "__main__": convert(sys.argv[1], sys.argv[2])
